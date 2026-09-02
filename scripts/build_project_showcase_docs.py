from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"

NAVY = "142B3D"
TEAL = "087F8C"
TEAL_DARK = "075F68"
TEAL_LIGHT = "EAF5F6"
INK = "17242D"
MUTED = "5E6C76"
LINE = "C8D1D6"
CANVAS = "F4F6F7"
WHITE = "FFFFFF"
GOLD = "8A6500"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float = 11, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 9, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document, running_label: str) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, TEAL_DARK, 18, 10),
        ("Heading 2", 13, TEAL_DARK, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run(running_label.upper()), 8.5, MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_base = max(existing_abs, default=0) + 1
    num_base = max(existing_num, default=0) + 1
    bullet_num = add_numbering_definition(numbering, abstract_base, num_base, bullet=True)
    decimal_num = add_numbering_definition(numbering, abstract_base + 1, num_base + 1, bullet=False)
    return bullet_num, decimal_num


def add_numbering_definition(numbering, abstract_id: int, num_id: int, bullet: bool) -> int:
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), 11, INK, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), 11, INK)
    else:
        set_run_font(p.add_run(text), 11, INK)


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str, metadata: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(kicker.upper()), 9, TEAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(title), 27, NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(subtitle), 13, MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run(metadata), 9.5, MUTED, italic=True)


def add_callout(doc: Document, label: str, text: str, fill: str = TEAL_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(label.upper()), 8.5, TEAL_DARK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), 11, INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), 9.5, WHITE, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), 9.5, INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_q_and_a(doc: Document, question: str, answer: str, bullet_num: int, follow_up: list[str] | None = None) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(question), 13, TEAL_DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Suggested answer: "), 11, NAVY, bold=True)
    set_run_font(p.add_run(answer), 11, INK)
    for item in follow_up or []:
        add_list_item(doc, item, bullet_num)


def build_showcase() -> Path:
    doc = Document()
    bullet_num, decimal_num = configure_document(doc, "Data Explorer | Project showcase")
    add_title_block(
        doc,
        "Project showcase guide",
        "Data Explorer",
        "A practical demonstration plan for governed enterprise RAG, publishing, and observability",
        "Prepared from the implemented repository and verified local workflows | 2 September 2026",
    )
    add_callout(
        doc,
        "Project in one sentence",
        "Data Explorer is a governed enterprise knowledge platform that retrieves only authorized evidence, produces citation-backed answers, applies independent approvals to high-risk actions, and exposes sanitized operational telemetry.",
    )

    doc.add_heading("How to present the project", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("Lead with the business risk, then prove the controls in the product. "), 11, NAVY, bold=True)
    set_run_font(p.add_run("Do not begin with a library list. Show that the system prevents unauthorized retrieval, unsupported answers, self-approval, unsafe SQL, and untraceable publishing."), 11, INK)

    doc.add_heading("90-second opening", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    set_run_font(
        p.add_run(
            '"I built Data Explorer to solve a common enterprise AI problem: a model can answer quickly, but the business still needs authorization, traceability, and safe publishing. The platform combines tenant-aware retrieval, citations, provider policy routing, structured-data approvals, document generation, evaluation, and a separate observability console. The key design choice is that identity and policy stay on the server, evidence is filtered before it reaches the model, and high-risk actions require independent approval."'
        ),
        11,
        INK,
        italic=True,
    )

    doc.add_heading("Architecture to explain", level=1)
    add_simple_table(
        doc,
        ["Layer", "Implemented responsibility", "Why it matters"],
        [
            ["Experience", "Streamlit workspace plus separate admin observability UI", "Keeps employee actions separate from privileged telemetry."],
            ["API and identity", "FastAPI, development or JWT authentication, tenant and group context", "Identity claims are resolved by the server and are not editable in the workspace."],
            ["Retrieval", "Chunking, embeddings, dense and sparse ranking, RRF, trust weighting, reranking", "Balances semantic matching with exact business terms while preserving ACL filters."],
            ["Generation", "Ollama, OpenAI, and Anthropic adapters behind a policy router", "Provider selection follows classification and deployment policy instead of prompt logic."],
            ["Governance", "Guardrails, citation checks, redaction, rate and token limits, approval workflows", "The system fails closed when evidence, permission, or policy is insufficient."],
            ["Operations", "Audit events, model traces, Prometheus metrics, Terraform and Cloud Build", "Supports production accountability without retaining prompt or response bodies in standard telemetry."],
        ],
        [1440, 4320, 3600],
    )

    doc.add_page_break()
    doc.add_heading("12-minute demonstration runbook", level=1)
    demo_steps = [
        "Open the workspace and admin panel. Show that workspace identity is server-derived and that the admin surface is separately authorized.",
        "Knowledge: ingest a short internal policy with a distinctive amount, approver, and schedule. Point out classification, trust tier, version, and allowed groups.",
        "Ask: query the distinctive facts. Show the grounded answer and citation, then explain that access filtering happened before model generation.",
        "Publish: create a DOCX approval request. Use a second reviewer identity, approve it, render it, and show the output hash and source register.",
        "Structured data: explain the propose, independently approve, and execute lifecycle. In local development, show the fail-closed response when no schema or executor is configured.",
        "Evaluate: run two golden questions. Show grounded count, citation count, provider, confidence, and latency rather than relying on anecdotal quality.",
        "Admin: show model traces, policy events, user activity, token totals, estimated cost, and grounding rate. Emphasize that raw prompts and answers are excluded from standard telemetry.",
        "Close on production architecture: JWT, PostgreSQL, Redis, Qdrant, versioned object storage, private networking, Terraform, monitoring, and immutable deployment images.",
    ]
    for item in demo_steps:
        add_list_item(doc, item, decimal_num)

    doc.add_heading("Suggested demonstration document", level=2)
    add_callout(
        doc,
        "Copy into Knowledge",
        "Regional procurement reviews occur every Tuesday. Purchase requests above AUD 12,750 require approval from the Procurement Director. Approved requests are released in the Friday supplier batch.",
        fill=CANVAS,
    )
    doc.add_heading("Questions to ask during the demo", level=2)
    for item in (
        "What purchase amount requires Procurement Director approval?",
        "When are approved purchase requests released?",
        "Who can retrieve this policy if its allowed group is finance?",
        "What should the system return when the caller has no authorized evidence?",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("What to point out in each workflow", level=1)
    workflow_rows = [
        ["Ask", "Citation-backed answers, retrieval confidence, refusal when evidence is insufficient", "Grounding is a product contract, not a prompt suggestion."],
        ["Knowledge", "Tenant, groups, classification, trust tier, version, checksum", "Governance metadata travels with every chunk."],
        ["Publish", "Typed ArtifactSpec, independent decision, source IDs, SHA-256", "The model cannot publish directly."],
        ["Structured data", "Schema policy, SQL AST validation, approver separation, read-only execution", "Generated SQL is treated as untrusted input."],
        ["Evaluate", "Golden questions, grounding, citations, confidence, provider, latency", "Quality is measured before release."],
        ["Admin", "Sanitized traces, audit events, users, tokens, cost, latency", "Observability supports accountability without storing business content."],
    ]
    add_simple_table(doc, ["Workflow", "Show", "Explain"], workflow_rows, [1440, 3600, 4320])

    doc.add_page_break()
    doc.add_heading("Technical depth to have ready", level=1)
    doc.add_heading("Retrieval and answer quality", level=2)
    for item in (
        "Dense similarity handles semantic matches; sparse BM25-style scoring preserves exact terms and identifiers.",
        "Reciprocal-rank fusion combines rankings; trust weighting favors authoritative content; lexical or FlashRank reranking refines final evidence.",
        "Expired sources and unauthorized chunks are removed before scoring and before model context assembly.",
        "Grounded responses require citations. If evidence is missing or confidence is too low, the service returns a calibrated limitation.",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("Security and governance", level=2)
    for item in (
        "Tenant and group authorization is enforced at ingestion, retrieval, SQL, artifacts, and observability boundaries.",
        "High-confidence prompt-injection patterns are blocked in both documents and questions.",
        "Provider routing considers classification; confidential content cannot silently fall back to an ineligible external provider.",
        "SQL is parsed with SQLGlot and limited to one bounded SELECT over allowlisted tables, columns, and functions.",
        "Publishing and SQL decisions require a different authorized user from the requester.",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("Production readiness", level=2)
    for item in (
        "Development defaults use memory and local Ollama for a low-friction vertical slice.",
        "Production refuses to start without JWT, PostgreSQL, Redis, Qdrant, and governed artifact storage configuration.",
        "Terraform provisions private networking, Cloud Run services and jobs, Artifact Registry, Cloud SQL, Memorystore, KMS, Secret Manager, monitoring, and budgets.",
        "Cloud Build produces immutable commit-tagged images; automated tests cover APIs, retrieval, policies, artifacts, observability, and UI rendering.",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("Known limitations to state confidently", level=1)
    add_simple_table(
        doc,
        ["Current limitation", "How to explain it", "Production path"],
        [
            ["Structured data is not configured in the default local server", "The UI deliberately returns 403 instead of inventing a schema or executing arbitrary SQL.", "Provide approved SchemaPolicy objects and a PostgreSQL read-only executor."],
            ["In-memory state is process-local", "It is suitable for tests and local demos, not durable production operation.", "Use PostgreSQL repositories, Redis policies, and Qdrant vectors."],
            ["PPTX rendering requires an external Node toolchain", "The API exposes the workflow but checks dependency reality before rendering.", "Provide the governed artifact runtime paths in deployment configuration."],
            ["LLM quality varies by model", "The system measures grounding and citations and can route providers by policy.", "Maintain domain golden sets and approved model catalogs."],
        ],
        [2160, 3600, 3600],
    )

    doc.add_heading("Final showcase checklist", level=1)
    for item in (
        "Workspace and admin URLs load without console errors.",
        "A uniquely identifiable document indexes successfully.",
        "The answer quotes the correct facts and includes at least one citation.",
        "An unauthorized identity receives no evidence.",
        "The artifact requester cannot self-approve; a second user can approve and render.",
        "The rendered output exists and its SHA-256 matches the API response.",
        "Structured data fails closed until a schema and executor are configured.",
        "Evaluation results and sanitized telemetry appear in the admin panel.",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("Close in one sentence", level=2)
    add_callout(
        doc,
        "Closing line",
        "The project demonstrates that enterprise AI is not only answer generation; it is authorization, evidence, approval, measurable quality, and operational accountability working as one system.",
    )

    path = OUTPUT_DIR / "DataExplorer_Project_Showcase_Guide.docx"
    doc.core_properties.title = "Data Explorer Project Showcase Guide"
    doc.core_properties.subject = "Demonstration guide for the Data Explorer enterprise RAG project"
    doc.core_properties.author = "Data Explorer"
    doc.save(path)
    return path


def build_interview_guide() -> Path:
    doc = Document()
    bullet_num, decimal_num = configure_document(doc, "Data Explorer | Interview preparation")
    add_title_block(
        doc,
        "Interview preparation",
        "Data Explorer: Questions and Suggested Answers",
        "Project-specific talking points for architecture, security, RAG, publishing, evaluation, and delivery",
        "Use the answers as adaptable prompts, not a script to memorize | 2 September 2026",
    )
    add_callout(
        doc,
        "Answer pattern",
        "Start with the decision or outcome, explain the design choice, name the control or tradeoff, and finish with how you verified it.",
    )

    doc.add_heading("Your opening answers", level=1)
    doc.add_heading("30-second project summary", level=2)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "Data Explorer is a governed enterprise RAG platform. It lets employees ingest authorized knowledge, ask citation-backed questions, create approved business artifacts, propose safe read-only SQL, run evaluations, and inspect sanitized telemetry. I designed the system so tenant and group authorization happens before retrieval, identity stays server-owned, and high-risk actions require independent approval."
        ),
        11,
        INK,
    )
    doc.add_heading("Two-minute walkthrough", level=2)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "The platform has a Streamlit workspace and a separately authorized observability console backed by FastAPI. Documents are validated, chunked, embedded, and indexed with tenant, ACL, classification, trust, version, and checksum metadata. Queries run through authorization, guardrails, hybrid dense and sparse retrieval, reciprocal-rank fusion, reranking, and a relevance threshold. Only eligible evidence reaches the model, and grounded responses must include citations. Structured SQL and business publishing use typed proposals plus independent approval. Development uses local memory and Ollama; production configuration is designed for JWT, PostgreSQL, Redis, Qdrant, governed object storage, private GCP networking, Terraform, and immutable builds."
        ),
        11,
        INK,
    )

    doc.add_page_break()
    doc.add_heading("Project and architecture questions", level=1)
    add_q_and_a(
        doc,
        "1. What problem does Data Explorer solve?",
        "It solves the gap between a useful AI answer and an enterprise-safe business process. A generic chatbot may retrieve the wrong tenant's data, produce unsupported claims, or publish without review. Data Explorer combines retrieval, authorization, citations, approval, evaluation, and audit so the answer is useful and governable.",
        bullet_num,
        ["Business outcome: faster access to internal knowledge with traceability.", "Risk outcome: fail closed on missing evidence, permission, or configuration."],
    )
    add_q_and_a(
        doc,
        "2. Walk me through the request flow.",
        "FastAPI authenticates the caller and creates a server-owned access context. Guardrails and budget checks validate the question. The retriever embeds the query, applies tenant and group filters, combines dense and sparse ranking, reranks evidence, and enforces a minimum relevance threshold. A policy router selects an eligible model, the answer is checked for grounding and citations, and sanitized trace and audit metadata is recorded.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "3. Why did you use FastAPI and Streamlit?",
        "FastAPI provides typed request and response models, dependency injection, asynchronous endpoints, and straightforward policy boundaries. Streamlit made it possible to build and test the five workflows quickly. I kept policy in the API so the UI remains replaceable and cannot become the security boundary.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "4. How is the architecture modular?",
        "Model providers, retrievers, audit stores, trace stores, policy enforcement, SQL executors, proposal repositories, and artifact storage are behind interfaces or replaceable implementations. Local memory and Ollama support development, while production adapters can use Qdrant, PostgreSQL, Redis, managed model providers, and governed object storage without rewriting domain logic.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "5. What is the most important design decision?",
        "Authorization happens before evidence reaches the model. Filtering after generation is too late because unauthorized content has already crossed the model boundary. The same deny-by-default idea is repeated for SQL, publishing, observability, and provider routing.",
        bullet_num,
    )

    doc.add_heading("RAG and model questions", level=1)
    add_q_and_a(
        doc,
        "6. Why combine dense and sparse retrieval?",
        "Dense vectors capture semantic similarity, while sparse term scoring protects exact identifiers, policy terms, amounts, and acronyms. The system uses reciprocal-rank fusion to combine the rankings, then applies trust weighting and a configurable reranker. This is more robust than relying on a single similarity score.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "7. How do you reduce hallucinations?",
        "I constrain generation to authorized evidence, require citations for grounded answers, verify that citations map to retrieved chunks, and return a limitation when evidence is insufficient. The system also records grounding, citation count, confidence, and retries so quality can be measured rather than assumed.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "8. How do you choose a model provider?",
        "Provider routing is configuration and policy, not prompt logic. Routes carry allowed classifications and whether they are external. Confidential or restricted content cannot silently fall back to a managed external provider. Ollama provides a private local route, while OpenAI and Anthropic adapters support eligible managed use cases.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "9. How does document ingestion work?",
        "The API validates document content and policy, chunks it by bounded structure, obtains embeddings, and indexes each chunk with tenant, groups, classification, trust tier, version, dates, metadata, and a content checksum. Deduplication and expiry logic prevent stale or repeated content from dominating retrieval.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "10. What would you measure for retrieval quality?",
        "I would track recall at k, precision at k, MRR or nDCG, reranker lift, authoritative-source preference, ACL leakage, groundedness, citation precision and recall, refusal correctness, latency, and cost. The golden set must include clean questions, noisy documents, stale versions, cross-tenant attacks, prompt injection, and insufficient-evidence cases.",
        bullet_num,
    )

    doc.add_page_break()
    doc.add_heading("Security and governance questions", level=1)
    add_q_and_a(
        doc,
        "11. How do you enforce multi-tenancy?",
        "The access context is derived from development headers or validated JWT claims. Tenant and group filters are applied before retrieval, and tenant checks are repeated for artifacts, SQL proposals, audit data, and observability. Production repositories also carry tenant identifiers so isolation is not dependent on the UI or a single filter.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "12. How do you handle prompt injection?",
        "Document and query text are validated before retrieval or generation. High-confidence injection patterns are blocked, provider context is policy-controlled, tools are narrowly scoped, and output is checked for schema, citations, and sensitive text. Most importantly, retrieved text never gets authority to change system policy.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "13. Why require independent approval?",
        "Generated SQL and business artifacts can create material consequences. The requester cannot approve their own proposal. A separate user must have the correct approver group and belong to the same tenant. That separation of duties turns human review into an enforced workflow state rather than a checkbox in the UI.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "14. How is Text2SQL made safer?",
        "The model sees only an approved schema description. SQLGlot parses the result and permits exactly one SELECT statement over allowlisted tables, columns, and functions. Comments, wildcards, DDL, DML, and system catalogs are blocked. After independent approval, PostgreSQL execution uses a read-only transaction, tenant session context, timeout, plan-cost ceiling, and row cap.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "15. Why does Structured data currently return 403 locally?",
        "That is an intentional fail-closed development state. The default server has no approved SchemaPolicy or SQL executor, so it refuses to generate or run SQL. To enable the workflow, I would configure the allowed finance schema and connect the PostgreSQL read-only executor. I would not invent a schema just to make the demo appear successful.",
        bullet_num,
    )

    doc.add_heading("Publishing and observability questions", level=1)
    add_q_and_a(
        doc,
        "16. How does governed publishing work?",
        "The requester submits a typed ArtifactSpec containing audience, purpose, classification, sections, claims, and source IDs. A different authorized reviewer approves or rejects it. Only an approved artifact can render. The output retains approval metadata, includes provenance, and is hashed so the published bytes can be verified.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "17. What do you record for observability?",
        "The system records correlation ID, tenant, user, operation, provider, model, token counts, latency, cost estimate, grounding, citation count, and bounded reflection attempts. It deliberately avoids retaining raw prompts, answers, or evidence in standard telemetry. The admin UI is separately authorized and tenant-scoped.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "18. How would you investigate a bad answer?",
        "I would start with the correlation ID, confirm caller and tenant context, inspect the retrieved source IDs and scores, verify document version and trust tier, check the selected provider and policy route, review grounding and citation validation, and compare the question with the relevant golden-set cases. The audit trail lets me reconstruct the decision without storing sensitive prompt bodies in normal telemetry.",
        bullet_num,
    )

    doc.add_page_break()
    doc.add_heading("Testing, deployment, and tradeoff questions", level=1)
    add_q_and_a(
        doc,
        "19. How did you test the system?",
        "The repository includes tests for APIs, retrieval quality, tenant guardrails, providers, configuration, Qdrant policy filters, Redis limits, artifacts, observability, Text2SQL validation, services, and Streamlit rendering. I also ran live end-to-end checks that ingested a unique document, asked grounded questions, verified an unauthorized refusal, independently approved and rendered a DOCX, ran a golden evaluation, and inspected the admin UI.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "20. What is different between development and production?",
        "Development favors a low-friction vertical slice with in-memory repositories, local policy enforcement, an in-memory retriever, and Ollama. Production refuses to start unless JWT, PostgreSQL persistence, Redis policy enforcement, Qdrant retrieval, and governed artifact storage are configured. That prevents accidental deployment with development-grade controls.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "21. How would you deploy it on GCP?",
        "I would deploy the API, workspace UI, and admin UI as separate Cloud Run services, with internal ingress and a stricter audience for admin. Long-running ingestion, evaluation, and rendering would run as jobs. Terraform provisions private networking, Cloud SQL, Memorystore, Qdrant connectivity, KMS, Secret Manager, monitoring, budgets, Artifact Registry, and versioned storage. Cloud Build produces immutable commit-tagged images.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "22. What tradeoffs did you make?",
        "I chose Streamlit for delivery speed, accepting that fine-grained UI control is more limited than a custom frontend. I kept memory-backed implementations for local development, accepting that state is process-local. I used deterministic lexical reranking by default to avoid model downloads, while keeping FlashRank optional. The architecture preserves replacement points so these choices do not become domain dependencies.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "23. What was a challenging bug or lesson?",
        "A useful example was frontend contrast. Streamlit rendered grouped controls under version-specific DOM selectors and inherited dark-theme colors even on a light page. I inspected the actual rendered roles and computed styles, then added explicit primary, secondary, selected, unselected, focus, and disabled states. I verified the contrast ratios in the browser instead of assuming the CSS selector worked.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "24. How would you improve performance and cost?",
        "I would profile ingestion and query latency separately, batch embeddings, cache only when tenant and ACL context are part of the key, route simple questions to lower-cost approved models, keep retrieval limits bounded, and alert on token and latency budgets. I would evaluate any cache or model change against quality and isolation metrics before release.",
        bullet_num,
    )
    add_q_and_a(
        doc,
        "25. What would you build next?",
        "First I would configure a real approved finance schema and PostgreSQL executor so the Structured data workflow can run end to end. Then I would add durable repositories, production identity integration, richer artifact preview and accessibility QA, a larger domain golden set, pagination and alerting in observability, and deployment runbooks for recovery, deletion, and incident response.",
        bullet_num,
    )

    doc.add_heading("Behavioral framing", level=1)
    add_simple_table(
        doc,
        ["Prompt", "Evidence from this project", "Result to emphasize"],
        [
            ["Tell me about ambiguity", "Translated enterprise RAG diagrams and a broad plan into replaceable boundaries and an incremental local slice.", "Made risk and production requirements explicit instead of hiding them in prompts."],
            ["Tell me about a failure", "Structured data correctly returned 403 with no configured schema; contrast selectors initially missed Streamlit's rendered component.", "Diagnosed with evidence, preserved fail-closed behavior, and fixed the UI using computed styles."],
            ["Tell me about security", "Tenant prefilters, provider classification routes, SQL AST validation, independent approval, sanitized telemetry.", "Applied defense in depth rather than relying on one guardrail library."],
            ["Tell me about quality", "Golden evaluations, citation checks, retrieval tests, browser QA, artifact hashing.", "Defined measurable release evidence instead of subjective demos."],
        ],
        [1800, 4200, 3360],
    )

    doc.add_heading("Questions you can ask the interviewer", level=1)
    for item in (
        "How does your team define acceptable grounding and citation quality for production AI?",
        "Where is authorization enforced today: source system, retrieval layer, application layer, or all three?",
        "Which AI actions require human approval, and how is separation of duties audited?",
        "How do you balance prompt and response observability with data-minimization requirements?",
        "What is the release process for changing models, prompts, retrieval strategies, or safety policies?",
        "Which production failure mode concerns the team most: data leakage, unsupported answers, cost, latency, or operational ownership?",
    ):
        add_list_item(doc, item, bullet_num)

    doc.add_heading("Last-minute checklist", level=1)
    for item in (
        "State the business problem before naming technologies.",
        "Use one concrete workflow example with a distinctive amount and citation.",
        "Explain one fail-closed behavior and why it is intentional.",
        "Separate what is implemented locally from the production target architecture.",
        "Name a tradeoff, a measured verification step, and the next improvement.",
        "Keep first answers concise; offer deeper detail when the interviewer follows up.",
    ):
        add_list_item(doc, item, bullet_num)

    path = OUTPUT_DIR / "DataExplorer_Interview_Questions_and_Answers.docx"
    doc.core_properties.title = "Data Explorer Interview Questions and Suggested Answers"
    doc.core_properties.subject = "Interview preparation for the Data Explorer enterprise RAG project"
    doc.core_properties.author = "Data Explorer"
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_showcase(), build_interview_guide()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
