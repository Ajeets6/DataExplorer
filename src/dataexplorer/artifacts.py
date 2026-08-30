import asyncio
import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from dataexplorer.models import (
    AccessContext,
    ArtifactDraft,
    ArtifactSection,
    ArtifactSpec,
)


class ArtifactPolicyError(ValueError):
    pass


class ArtifactRenderError(RuntimeError):
    pass


class ArtifactRenderer(Protocol):
    kind: str

    async def render(
        self,
        draft: ArtifactDraft,
        output_directory: Path,
    ) -> tuple[Path, Path | None]: ...


class ArtifactPublisher(Protocol):
    async def publish(self, source: Path, object_name: str) -> str: ...


@dataclass(slots=True)
class InMemoryArtifactRepository:
    drafts: dict[str, ArtifactDraft] = field(default_factory=dict)
    _lock: asyncio.Lock = field(default_factory=asyncio.Lock)

    async def save(self, draft: ArtifactDraft) -> None:
        async with self._lock:
            self.drafts[draft.artifact_id] = draft

    async def get(self, artifact_id: str) -> ArtifactDraft:
        async with self._lock:
            draft = self.drafts.get(artifact_id)
        if draft is None:
            raise ArtifactPolicyError("artifact draft was not found")
        return draft


@dataclass(slots=True)
class ArtifactService:
    output_root: Path
    renderers: dict[str, ArtifactRenderer]
    repository: InMemoryArtifactRepository = field(
        default_factory=InMemoryArtifactRepository
    )
    approver_group: str = "content-approvers"
    publisher: ArtifactPublisher | None = None

    async def create(self, spec: ArtifactSpec, access: AccessContext) -> ArtifactDraft:
        draft = ArtifactDraft(
            spec=spec,
            tenant_id=access.tenant_id,
            requested_by=access.user_id,
        )
        await self.repository.save(draft)
        return draft

    async def decide(
        self,
        artifact_id: str,
        *,
        approved: bool,
        reason: str,
        access: AccessContext,
    ) -> ArtifactDraft:
        draft = await self.repository.get(artifact_id)
        self._same_tenant(draft, access)
        if self.approver_group not in access.groups:
            raise ArtifactPolicyError("caller is not an authorized content approver")
        if draft.requested_by == access.user_id:
            raise ArtifactPolicyError("requesters cannot approve their own artifact")
        if draft.status != "pending":
            raise ArtifactPolicyError("only pending artifacts can be reviewed")
        updated = draft.model_copy(
            update={
                "status": "approved" if approved else "rejected",
                "approved_by": access.user_id,
                "approval_reason": reason,
            }
        )
        await self.repository.save(updated)
        return updated

    async def render(self, artifact_id: str, access: AccessContext) -> ArtifactDraft:
        draft = await self.repository.get(artifact_id)
        self._same_tenant(draft, access)
        if draft.status != "approved":
            raise ArtifactPolicyError("artifact must be approved before rendering")
        renderer = self.renderers.get(draft.spec.kind)
        if renderer is None:
            raise ArtifactRenderError(f"no renderer is configured for {draft.spec.kind}")
        tenant_root = (self.output_root / draft.tenant_id / draft.artifact_id).resolve()
        root = self.output_root.resolve()
        if root not in tenant_root.parents:
            raise ArtifactRenderError("artifact output escaped the configured root")
        tenant_root.mkdir(parents=True, exist_ok=True)
        output_path, qa_manifest = await renderer.render(draft, tenant_root)
        digest = hashlib.sha256(output_path.read_bytes()).hexdigest()
        published_output = str(output_path)
        published_manifest = str(qa_manifest) if qa_manifest else None
        if self.publisher:
            prefix = f"{draft.tenant_id}/{draft.artifact_id}"
            published_output = await self.publisher.publish(
                output_path, f"{prefix}/{output_path.name}"
            )
            if qa_manifest:
                published_manifest = await self.publisher.publish(
                    qa_manifest, f"{prefix}/qa/{qa_manifest.name}"
                )
        updated = draft.model_copy(
            update={
                "status": "rendered",
                "output_path": published_output,
                "output_sha256": digest,
                "qa_manifest_path": published_manifest,
            }
        )
        await self.repository.save(updated)
        return updated

    @staticmethod
    def _same_tenant(draft: ArtifactDraft, access: AccessContext) -> None:
        if draft.tenant_id != access.tenant_id:
            raise ArtifactPolicyError("artifact belongs to another tenant")


@dataclass(slots=True)
class DocxRenderer:
    kind: str = "docx"

    async def render(
        self,
        draft: ArtifactDraft,
        output_directory: Path,
    ) -> tuple[Path, Path | None]:
        output = output_directory / draft.spec.filename
        await asyncio.to_thread(_write_docx, draft, output)
        return output, None


@dataclass(slots=True)
class GcsArtifactPublisher:
    bucket_name: str
    kms_key_name: str | None = None

    async def publish(self, source: Path, object_name: str) -> str:
        return await asyncio.to_thread(self._publish, source, object_name)

    def _publish(self, source: Path, object_name: str) -> str:
        from google.cloud import storage

        client = storage.Client()
        bucket = client.bucket(self.bucket_name)
        blob = bucket.blob(object_name, kms_key_name=self.kms_key_name)
        blob.upload_from_filename(source, if_generation_match=0)
        return f"gs://{self.bucket_name}/{object_name}"


@dataclass(slots=True)
class ArtifactToolPptxRenderer:
    node_executable: Path
    node_modules: Path
    script_path: Path
    kind: str = "pptx"

    async def render(
        self,
        draft: ArtifactDraft,
        output_directory: Path,
    ) -> tuple[Path, Path | None]:
        output = output_directory / draft.spec.filename
        spec_path = output_directory / "artifact-spec.json"
        qa_dir = output_directory / "qa"
        spec_path.write_text(draft.spec.model_dump_json(indent=2), encoding="utf-8")
        qa_dir.mkdir(exist_ok=True)
        environment = {
            **os.environ,
            "NODE_PATH": str(self.node_modules),
            "ARTIFACT_TOOL_NODE_MODULES": str(self.node_modules),
        }
        process = await asyncio.create_subprocess_exec(
            str(self.node_executable),
            str(self.script_path),
            str(spec_path),
            str(output),
            str(qa_dir),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=environment,
        )
        stdout, stderr = await process.communicate()
        if process.returncode != 0:
            raise ArtifactRenderError(
                f"PowerPoint renderer failed: {(stderr or stdout).decode(errors='replace')}"
            )
        manifest = qa_dir / "manifest.json"
        if not output.exists() or not manifest.exists():
            raise ArtifactRenderError("PowerPoint renderer did not produce required outputs")
        return output, manifest


def _write_docx(draft: ArtifactDraft, output: Path) -> None:
    spec = draft.spec
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    _configure_styles(document)
    _add_header_footer(section, spec.classification, draft.artifact_id)
    _add_memo_masthead(document, draft)
    for content in spec.sections:
        document.add_heading(content.title, level=1)
        if content.summary:
            paragraph = document.add_paragraph(content.summary)
            paragraph.style = document.styles["Normal"]
        for bullet in content.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        if content.table:
            _add_table(document, content)
        if content.chart:
            _add_chart_data_table(document, content)
        _add_source_line(document, content, spec)
    document.add_heading("Source register", level=1)
    for source in spec.sources:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{source.source_id}: {source.label} — {source.locator}")
    core = document.core_properties
    core.title = spec.title
    core.subject = spec.purpose
    core.author = "Data Explorer"
    core.comments = (
        f"Classification={spec.classification}; Template={spec.template_version}; "
        f"ApprovedBy={draft.approved_by}; ArtifactId={draft.artifact_id}"
    )
    document.save(output)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    bullet = styles["List Bullet"]
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167


def _add_header_footer(section, classification: str, artifact_id: str) -> None:
    header = section.header.paragraphs[0]
    header.text = f"DATA EXPLORER  |  {classification.upper()}"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer.paragraphs[0]
    footer.text = f"Generated artifact  |  {artifact_id}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("777777")


def _add_memo_masthead(document: Document, draft: ArtifactDraft) -> None:
    spec = draft.spec
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(spec.title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    subtitle = document.add_paragraph(spec.subtitle)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("444444")
    for label, value in (
        ("Audience", spec.audience),
        ("Purpose", spec.purpose),
        ("Classification", spec.classification.upper()),
        ("Approved by", draft.approved_by or "Pending"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)
    rule = document.add_paragraph()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "2E74B5")
    border.append(bottom)
    rule._p.get_or_add_pPr().append(border)


def _add_table(document: Document, section: ArtifactSection) -> None:
    table_spec = section.table
    assert table_spec is not None
    document.add_paragraph(table_spec.title, style="Heading 2")
    table = document.add_table(rows=1, cols=len(table_spec.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = _table_widths(table_spec.columns)
    for index, heading in enumerate(table_spec.columns):
        cell = table.rows[0].cells[index]
        cell.text = heading
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
    for values in table_spec.rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    _set_table_geometry(table, widths)


def _add_chart_data_table(document: Document, section: ArtifactSection) -> None:
    chart = section.chart
    assert chart is not None
    document.add_paragraph(f"{chart.title} — governed chart data", style="Heading 2")
    table = document.add_table(rows=1, cols=1 + len(chart.series))
    headings = ["Category", *[series.name for series in chart.series]]
    for index, heading in enumerate(headings):
        table.rows[0].cells[index].text = heading
        _shade_cell(table.rows[0].cells[index], "E8EEF5")
    for category_index, category in enumerate(chart.categories):
        row = table.add_row()
        row.cells[0].text = category
        for series_index, series in enumerate(chart.series, start=1):
            row.cells[series_index].text = str(series.values[category_index])
    widths = [2.5] + [4.0 / len(chart.series)] * len(chart.series)
    _set_table_geometry(table, widths)


def _add_source_line(document: Document, section: ArtifactSection, spec: ArtifactSpec) -> None:
    lookup = {source.source_id: source for source in spec.sources}
    labels = [lookup[source_id].label for source_id in sorted(section.source_ids)]
    paragraph = document.add_paragraph(f"Sources: {'; '.join(labels)}")
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("666666")


def _table_widths(columns: list[str]) -> list[float]:
    weights = [max(1.0, min(3.0, len(column) / 10)) for column in columns]
    total = sum(weights)
    return [6.5 * weight / total for weight in weights]


def _set_table_geometry(table, widths: list[float]) -> None:
    table_element = table._tbl
    properties = table_element.tblPr
    width_element = properties.first_child_found_in("w:tblW")
    width_element.set(qn("w:type"), "dxa")
    width_element.set(qn("w:w"), "9360")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        dxa = round(width * 1440)
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(dxa))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            dxa = round(widths[index] * 1440)
            cell.width = Inches(widths[index])
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(dxa))


def _shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)
